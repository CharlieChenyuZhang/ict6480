from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import ListFlowable, ListItem, PageBreak, Paragraph, SimpleDocTemplate, Spacer

ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / 'output/doc/Course Outline Finalized Contents.md'
DOCX_OUT = ROOT / 'output/doc/Course Outline Finalized Contents.docx'
PDF_OUT = ROOT / 'output/doc/Course Outline Finalized Contents.pdf'


def clean_inline(text: str) -> str:
    text = text.strip()
    if text.startswith('**') and text.endswith('**'):
        text = text[2:-2]
    return text.replace('**', '')


def parse_blocks(text: str):
    blocks = []
    lines = text.splitlines()
    buffer = []

    def flush_paragraph():
        nonlocal buffer
        if buffer:
            paragraph = ' '.join(part.strip() for part in buffer if part.strip())
            if paragraph:
                blocks.append(('paragraph', paragraph))
        buffer = []

    for line in lines:
        raw = line.rstrip()
        stripped = raw.strip()
        if not stripped:
            flush_paragraph()
            continue
        if stripped.startswith('### '):
            flush_paragraph()
            blocks.append(('heading3', stripped[4:].strip()))
            continue
        if stripped.startswith('## '):
            flush_paragraph()
            blocks.append(('heading2', stripped[3:].strip()))
            continue
        if stripped.startswith('# '):
            flush_paragraph()
            blocks.append(('heading1', stripped[2:].strip()))
            continue
        if stripped.startswith('- '):
            flush_paragraph()
            blocks.append(('bullet', stripped[2:].strip()))
            continue
        if stripped.startswith('**') and stripped.endswith('**'):
            flush_paragraph()
            blocks.append(('label', clean_inline(stripped)))
            continue
        buffer.append(stripped)

    flush_paragraph()
    return blocks


def add_page_number(paragraph):
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def generate_docx(blocks):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = doc.styles
    styles['Normal'].font.name = 'Aptos'
    styles['Normal'].font.size = Pt(10.5)
    styles['Heading 1'].font.name = 'Aptos Display'
    styles['Heading 1'].font.size = Pt(16)
    styles['Heading 1'].font.bold = True
    styles['Heading 2'].font.name = 'Aptos Display'
    styles['Heading 2'].font.size = Pt(12.5)
    styles['Heading 2'].font.bold = True
    styles['Heading 3'].font.name = 'Aptos Display'
    styles['Heading 3'].font.size = Pt(11.5)
    styles['Heading 3'].font.bold = True

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('COURSE OUTLINE')
    run.bold = True
    run.font.size = Pt(18)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('AWS Cloud Architecting')
    run.bold = True
    run.font.size = Pt(14)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run('Prepared from course template, prior ITC6480 content, meeting notes, and AWS Academy Cloud Architecting materials.').italic = True

    doc.add_paragraph()

    for kind, value in blocks:
        if kind == 'heading1':
            continue
        if kind == 'heading2':
            doc.add_heading(value, level=1)
        elif kind == 'heading3':
            doc.add_heading(value, level=2)
        elif kind == 'label':
            p = doc.add_paragraph()
            r = p.add_run(value)
            r.bold = True
        elif kind == 'bullet':
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(clean_inline(value))
        elif kind == 'paragraph':
            doc.add_paragraph(clean_inline(value))

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run('ITC6480 / AWS Cloud Architecting - ')
    add_page_number(footer)

    doc.save(DOCX_OUT)


def xml_escape(text: str) -> str:
    return (
        text.replace('&', '&amp;')
        .replace('<', '&lt;')
        .replace('>', '&gt;')
    )


def generate_pdf(blocks):
    doc = SimpleDocTemplate(
        str(PDF_OUT),
        pagesize=LETTER,
        leftMargin=0.8 * inch,
        rightMargin=0.8 * inch,
        topMargin=0.7 * inch,
        bottomMargin=0.7 * inch,
        title='Course Outline - AWS Cloud Architecting',
    )
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'TitleCenter',
        parent=styles['Title'],
        alignment=TA_CENTER,
        fontName='Helvetica-Bold',
        fontSize=18,
        leading=22,
        textColor=colors.black,
        spaceAfter=6,
    )
    subtitle_style = ParagraphStyle(
        'SubtitleCenter',
        parent=styles['Heading2'],
        alignment=TA_CENTER,
        fontName='Helvetica-Bold',
        fontSize=13,
        leading=16,
        spaceAfter=6,
    )
    note_style = ParagraphStyle(
        'NoteCenter',
        parent=styles['BodyText'],
        alignment=TA_CENTER,
        fontName='Helvetica-Oblique',
        fontSize=9.5,
        leading=12,
        spaceAfter=14,
    )
    h1_style = ParagraphStyle(
        'H1',
        parent=styles['Heading1'],
        fontName='Helvetica-Bold',
        fontSize=13,
        leading=16,
        spaceBefore=10,
        spaceAfter=6,
        textColor=colors.HexColor('#1F3A5F'),
    )
    h2_style = ParagraphStyle(
        'H2',
        parent=styles['Heading2'],
        fontName='Helvetica-Bold',
        fontSize=11.5,
        leading=14,
        spaceBefore=8,
        spaceAfter=4,
        textColor=colors.HexColor('#2A5D84'),
    )
    label_style = ParagraphStyle(
        'Label',
        parent=styles['BodyText'],
        fontName='Helvetica-Bold',
        fontSize=10,
        leading=12,
        spaceBefore=4,
        spaceAfter=2,
    )
    body_style = ParagraphStyle(
        'Body',
        parent=styles['BodyText'],
        fontName='Helvetica',
        fontSize=9.5,
        leading=12,
        spaceAfter=6,
    )
    bullet_style = ParagraphStyle(
        'BulletBody',
        parent=body_style,
        leftIndent=12,
        firstLineIndent=0,
        spaceAfter=2,
    )

    story = [
        Paragraph('COURSE OUTLINE', title_style),
        Paragraph('AWS Cloud Architecting', subtitle_style),
        Paragraph('Prepared from course template, prior ITC6480 content, meeting notes, and AWS Academy Cloud Architecting materials.', note_style),
        Spacer(1, 0.08 * inch),
    ]

    bullet_buffer = []

    def flush_bullets():
        nonlocal bullet_buffer
        if bullet_buffer:
            items = [ListItem(Paragraph(xml_escape(clean_inline(item)), bullet_style)) for item in bullet_buffer]
            story.append(ListFlowable(items, bulletType='bullet', start='circle', leftIndent=18))
            bullet_buffer = []

    for kind, value in blocks:
        if kind == 'heading1':
            continue
        if kind in {'heading2', 'heading3', 'label', 'paragraph'}:
            flush_bullets()
        if kind == 'heading2':
            story.append(Paragraph(xml_escape(value), h1_style))
        elif kind == 'heading3':
            story.append(Paragraph(xml_escape(value), h2_style))
        elif kind == 'label':
            story.append(Paragraph(xml_escape(value), label_style))
        elif kind == 'paragraph':
            story.append(Paragraph(xml_escape(clean_inline(value)), body_style))
        elif kind == 'bullet':
            bullet_buffer.append(value)

    flush_bullets()

    def draw_footer(canvas, _doc):
        canvas.saveState()
        canvas.setFont('Helvetica', 9)
        canvas.setFillColor(colors.grey)
        canvas.drawCentredString(LETTER[0] / 2.0, 0.45 * inch, f'ITC6480 / AWS Cloud Architecting - Page {canvas.getPageNumber()}')
        canvas.restoreState()

    doc.build(story, onFirstPage=draw_footer, onLaterPages=draw_footer)


def main():
    text = SOURCE.read_text()
    blocks = parse_blocks(text)
    generate_docx(blocks)
    generate_pdf(blocks)
    print(DOCX_OUT)
    print(PDF_OUT)


if __name__ == '__main__':
    main()
